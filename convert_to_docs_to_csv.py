#!pip install python-docx
#!pip install pandas
#!pip install numpy
#!pip install Apache-airflow
import docx
import pandas as pd
import numpy as np
from collections import defaultdict
import copy
from airflow import DAG
from airflow.operators.bash import BashOperator
from airflow.operators.python import PythonOperator
import datetime as dt
from pathlib import Path


#templ = docx.Document('/Users/pn_jh/Downloads/기상청27_관광코스별_관광지_상세_날씨_조회서비스_오픈API활용가이드 (1).docx')
#templ = templ.tables
input_path='/Users/pn_jh/Downloads/기상청27_관광코스별_관광지_상세_날씨_조회서비스_오픈API활용가이드 (1).docx'

#templ[0].rows[0].cells[0].text
def _read_docx_and_make_dataframe(input_path):
    templ = docx.Document(input_path)
    templ = templ.tables
    df=pd.DataFrame(columns=['courseId','courseName','Tourist Spot'])
    # 14번째 행부터 96번째 행까지가 관광지 정보 그리고 97번쨰부터는 지역구정보
    for x in range(14,96+1):
        #print(x)
        for table_rows in templ[x].rows:
            temp=dict()
            i=0
            for table_cells in table_rows.cells:
                #print(table_cells.text,x)
                if i==0:
                    temp['courseId']=table_cells.text
                elif i==1:
                    temp['courseName']=table_cells.text
                else:
                    temp['Tourist Spot']=table_cells.text
                i+=1
            df=df.append(temp,ignore_index=True)
    df=copy.deepcopy(df.iloc[1:,:].reset_index(drop=True))
    df.to_csv('/Users/pn_jh/Desktop/DockerProjects/weather/course_information.csv',encoding='utf-8-sig')
    return 
_read_docx_and_make_dataframe(input_path)
